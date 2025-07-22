import csv
import io
import time
import uuid
import fitz  
import docx
import pandas as pd
from typing import Optional, Dict, List
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document


class EmbeddingService:
    def __init__(self, client, embedding_model):
        self.client = client
        self.embedding_model = embedding_model
        self.text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)

    def get_collection_name(self, user_id: str, agent_id: str) -> str:
        return f"user_{user_id}_agent_{agent_id}"

    async def create_collection(self, user_id: str, agent_id: str) -> None:
        collection_name = self.get_collection_name(user_id, agent_id)
        try:
            self.client.get_collection(collection_name)
        except:
            self.client.create_collection(
                collection_name=collection_name,
                vectors_config={"size": 3072, "distance": "Cosine"}
            )

    def _extract_text(self, file_bytes: bytes, file_type: str) -> str:
        if file_type == "application/pdf":
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            return "\n".join([page.get_text() for page in doc])
        elif file_type == "text/plain":
            return file_bytes.decode("utf-8")
        elif file_type == "text/csv":
            decoded = file_bytes.decode("utf-8")
            reader = csv.reader(io.StringIO(decoded))
            return "\n".join([", ".join(row) for row in reader])
        elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            doc = docx.Document(io.BytesIO(file_bytes))
            return "\n".join([p.text for p in doc.paragraphs])
        elif file_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            df = pd.read_excel(io.BytesIO(file_bytes))
            return df.to_csv(index=False)
        elif file_type == "text/markdown":
            return file_bytes.decode("utf-8")
        else:
            return file_bytes.decode("utf-8", errors="ignore")

    async def embed_uploaded_document(
        self,
        file_bytes: bytes,
        file_type: str,
        filename: str,
        user_id: str,
        agent_id: str,
        file_id: uuid.UUID,
        custom_metadata: Optional[Dict] = None,
    ) -> Dict:
        collection_name = self.get_collection_name(user_id, agent_id)
        await self.create_collection(user_id, agent_id)

        text = self._extract_text(file_bytes, file_type)
        documents = [Document(page_content=text, metadata={"filename": filename})]
        chunks = self.text_splitter.split_documents(documents)

        texts = [chunk.page_content for chunk in chunks]
        embeddings = await self.embedding_model.aembed_documents(texts)

        timestamp = int(time.time())
        points = []
        for chunk, embedding in zip(chunks, embeddings):
            metadata = {
                "user_id": user_id,
                "agent_id": agent_id,
                "upload_timestamp": timestamp,
                **chunk.metadata
            }

            if custom_metadata:
                metadata.update(custom_metadata)

            points.append({
                "id": str(uuid.uuid4()),
                "vector": embedding,
                "payload": {
                    "text": chunk.page_content,
                    "metadata": metadata
                }
            })

        self.client.upsert(collection_name=collection_name, points=points)

        return {
            "status": "success",
            "chunks_processed": len(points),
            "collection": collection_name,
            "document_id": str(file_id)
        }


    async def search_for_context(
        self,
        input: str,
        agent_id: str,
        user_id: str,
        tok_k: int = 4
    ) -> List[Document]:
        collection_name = self.get_collection_name(user_id=user_id, agent_id=agent_id)
        query_embedding = await self.embedding_model.aembed_query(input)

        search_results = self.client.search(
            collection_name=collection_name,
            query_vector=query_embedding,
            limit=tok_k,
            with_payload=True
        )

        if not search_results:
            return None

        docs =  [
            Document(
                page_content=item.payload["text"]
            ) for item in search_results
        ]

        return  "\n\n".join([doc.page_content for doc in docs])


    def scroll(self, user_id: str, agent_id: str):
        results, _ = self.client.scroll(
            collection_name=f"user_{user_id}_agent_{agent_id}",
            limit=10,
            with_payload=True
        )

        for point in results:
            print(point.payload)